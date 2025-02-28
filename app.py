import os
import re
import cv2
import pytesseract
import pdfplumber
from flask import Flask, request, render_template, redirect, url_for, session
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

app = Flask(__name__)
app.secret_key = "supersecretkey"  # Required for session handling

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Path to Tesseract OCR
pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

# Function to clean extracted text
def clean_extracted_text(text):
    if not text:
        return ""
    text = re.sub(r'\s*\|\s*', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text.lower()

# Extract hyperlinks from PDFs
def extract_links_from_pdf(pdf_path):
    links = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            if page.annots:
                for annotation in page.annots:
                    if "/A" in annotation and "/URI" in annotation["/A"]:
                        links.append(annotation["/A"]["/URI"])
    return links

# Extract hyperlinks from DOCX files
def extract_links_from_docx(docx_path):
    doc = Document(docx_path)
    links = []
    for rel in doc.part.rels.values():
        if rel.reltype == RELATIONSHIP_TYPE.HYPERLINK:
            links.append(rel.target_ref)
    return links

# Function to extract text from different file types and fetch links
def extract_text(file_path, file_ext):
    text, links = "", []
    try:
        if file_ext == "pdf":
            reader = PdfReader(file_path)
            text = "".join([page.extract_text() or "" for page in reader.pages])
            links = extract_links_from_pdf(file_path)
        elif file_ext in ["docx", "doc"]:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            links = extract_links_from_docx(file_path)
        elif file_ext in ["png", "jpg", "jpeg"]:
            img = cv2.imread(file_path)
            text = pytesseract.image_to_string(img)
        else:
            return "Unsupported file type", []
    except Exception as e:
        return f"Error extracting text: {str(e)}", []
    
    return clean_extracted_text(text), links

# Function to check ATS compliance
def check_ats_compliance(text):
    if len(text) < 200:
        return "Resume too short. Consider adding more content."
    if len(text.split()) < 100:
        return "Low word count detected. Expand on your experiences."
    return "ATS-friendly format detected."

# Function to extract keywords and determine selection status
def extract_and_highlight_keywords(text, links):
    keywords = [
        'python', 'java', 'c++', 'c#', 'javascript', 'sql', 'aws', 'azure', 'tensorflow',
        'keras', 'pandas', 'numpy', 'ml', 'ai', 'deep learning', 'data science'
    ]
    matched_keywords = [word for word in keywords if word in text]
    ats_friendly = check_ats_compliance(text)
    selection_status = "Selected" if matched_keywords else "Not Selected"

    linkedin_url = next((link for link in links if "linkedin.com/in" in link), "Not Provided")
    github_url = next((link for link in links if "github.com" in link), "Not Provided")

    return selection_status, matched_keywords, ats_friendly, linkedin_url, github_url

# Function to check job role match score
def job_role_matching(resume_text, job_desc):
    if not job_desc.strip():
        return 0  # If no job description provided, return 0% match
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform([resume_text, job_desc])
    similarity_score = cosine_similarity(tfidf_matrix[0], tfidf_matrix[1])[0][0] * 100
    return round(similarity_score, 2)

# Home Page Route
@app.route("/", methods=["GET"])
def home():
    return render_template("Home.html")

# Resume Selection Route
@app.route("/selection", methods=["GET", "POST"])
def selection():
    if request.method == "POST":
        if 'file' not in request.files:
            session['result'] = {"error": "No file uploaded"}
            return redirect(url_for("selection"))

        file = request.files['file']
        job_desc = request.form.get("job_desc", "")

        if file.filename == "":
            session['result'] = {"error": "No file selected"}
            return redirect(url_for("selection"))

        filename = secure_filename(file.filename)
        file_ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

        if file_ext not in ["pdf", "docx", "doc", "png", "jpg", "jpeg"]:
            session['result'] = {"error": "Unsupported file type"}
            return redirect(url_for("selection"))

        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        try:
            resume_text, links = extract_text(file_path, file_ext)
            if "Unsupported file type" in resume_text or "Error extracting text" in resume_text:
                session['result'] = {"error": resume_text}
                return redirect(url_for("selection"))

            if not resume_text.strip():
                session['result'] = {"error": "No text extracted from the file. Please use a different format."}
                return redirect(url_for("selection"))

            selection_status, matched_keywords, ats_friendly, linkedin_url, github_url = extract_and_highlight_keywords(resume_text, links)
            similarity_score = job_role_matching(resume_text, job_desc)

            session['result'] = {
                "Selection Status": selection_status,
                "Matched Keywords": matched_keywords,
                "ATS Compliance": ats_friendly,
                "Job Match Score (%)": similarity_score,
                "LinkedIn Profile": f"<a href='{linkedin_url}' target='_blank'>{linkedin_url}</a>" if linkedin_url != "Not Provided" else "Not Provided",
                "GitHub Profile": f"<a href='{github_url}' target='_blank'>{github_url}</a>" if github_url != "Not Provided" else "Not Provided"
            }

            return redirect(url_for("selection"))

        except Exception as e:
            session['result'] = {"error": f"Unexpected error: {str(e)}"}
            return redirect(url_for("selection"))

    result = session.pop('result', None)  # Retrieve and remove result from session
    return render_template("Selection.html", result=result)

# Run the Flask app
if __name__ == "__main__":
    app.run(debug=True)
